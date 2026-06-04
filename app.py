import os
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

import streamlit as st
import streamlit.components.v1 as components
import base64
import io
import time
import json
import tempfile
import math
import shutil
import subprocess
import logging
import sys
import traceback
import platform
from datetime import datetime
from pathlib import Path

# ── Tab Recorder Component ────────────────────────────────────────────────────
# This component captures browser tab audio using getDisplayMedia
COMPONENT_DIR = Path(__file__).parent.absolute() / "meeting_recorder"
COMPONENT_DIR.mkdir(exist_ok=True)

# index.html content is already written by the setup step, 
# but we keep a reference to the HTML for consistency.
tab_recorder_component = components.declare_component("tab_recorder", path=str(COMPONENT_DIR))

def render_tab_recorder():
    """Renders the custom tab audio recorder with a fallback message."""
    try:
        data = tab_recorder_component(key="tab_audio_recorder")
        return data
    except Exception as e:
        st.error(f"Failed to load the meeting recorder component: {e}")
        st.info("Try refreshing the page. Ensure you are using a modern browser (Chrome, Edge).")
        return None


# ── torchaudio Monkeypatch ────────────────────────────────────────────────────
# Resolve AttributeError: module 'torchaudio' has no attribute 'set_audio_backend'
# This occurs because pyannote.audio (including some 4.x versions) calls a 
# removed API in torchaudio 2.1.0+.
try:
    import torchaudio
    if not hasattr(torchaudio, "set_audio_backend"):
        torchaudio.set_audio_backend = lambda x: None
        logging.info("Applied monkeypatch for torchaudio.set_audio_backend")
except Exception as e:
    logging.debug(f"Could not load or patch torchaudio: {e}")

# ── Logging Configuration ─────────────────────────────────────────────────────
LOG_DIR = Path("logs")
LOG_DIR.mkdir(exist_ok=True)
LOG_FILE = LOG_DIR / "application.log"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("MeetingMind")

# ── System Diagnostics ────────────────────────────────────────────────────────
def check_audio_ai_environment():
    """Exhaustive check of the ML and audio processing environment."""
    import torch
    import torchaudio
    try:
        import pyannote.audio
        pyannote_ver = pyannote.audio.__version__
    except ImportError:
        pyannote_ver = "NOT_INSTALLED"
    
    try:
        import faster_whisper
        whisper_ver = faster_whisper.__version__
    except ImportError:
        whisper_ver = "NOT_INSTALLED"

    try:
        import torchcodec
        codec_ver = getattr(torchcodec, "__version__", "Installed")
    except ImportError:
        codec_ver = "MISSING"

    results = {
        "Python": sys.version.split()[0],
        "Platform": platform.platform(),
        "Torch": torch.__version__,
        "TorchAudio": torchaudio.__version__,
        "PyAnnote.Audio": pyannote_ver,
        "Faster-Whisper": whisper_ver,
        "TorchCodec": codec_ver,
        "FFmpeg Found": shutil.which("ffmpeg") is not None,
        "CUDA Available": torch.cuda.is_available(),
    }
    
    # Check for missing FFmpeg libraries if on Linux/Streamlit Cloud
    if platform.system() == "Linux":
        try:
            # Check for libavutil
            res = subprocess.run(["ldconfig", "-p"], capture_output=True, text=True)
            results["libavutil"] = "Found" if "libavutil" in res.stdout else "Missing"
        except:
            results["libavutil"] = "Unknown (ldconfig failed)"

    return results

def get_system_diagnostics():
    import torch
    diag = {
        "timestamp": datetime.now().isoformat(),
        "env": check_audio_ai_environment(),
    }
    
    # Check HF Token
    hf_token = st.secrets.get("HF_TOKEN", "") or os.environ.get("HF_TOKEN", "")
    diag["hf_token_present"] = bool(hf_token)
    diag["hf_token_valid"] = "Not Verified"
    
    if hf_token:
        try:
            from huggingface_hub import HfApi
            api = HfApi(token=hf_token)
            api.whoami()
            diag["hf_token_valid"] = True
        except Exception as e:
            diag["hf_token_valid"] = False
            diag["hf_error"] = str(e)
            
    return diag

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="MeetingMind",
    page_icon="🎙",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Custom CSS ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display&family=DM+Sans:wght@300;400;500;600&display=swap');

html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
}

/* Header */
.main-header {
    background: linear-gradient(135deg, #0f1117 0%, #1a1f2e 50%, #0f1117 100%);
    border: 1px solid #2a3050;
    border-radius: 16px;
    padding: 28px 32px;
    margin-bottom: 24px;
    position: relative;
    overflow: hidden;
}
.main-header::before {
    content: '';
    position: absolute;
    top: -50%;
    left: -50%;
    width: 200%;
    height: 200%;
    background: radial-gradient(ellipse at 30% 50%, rgba(99,102,241,0.08) 0%, transparent 60%),
                radial-gradient(ellipse at 70% 50%, rgba(16,185,129,0.06) 0%, transparent 60%);
    pointer-events: none;
}
.main-header h1 {
    font-family: 'DM Serif Display', serif;
    font-size: 2.2rem;
    color: #f1f5f9;
    margin: 0;
    letter-spacing: -0.5px;
}
.main-header p {
    color: #94a3b8;
    margin: 6px 0 0;
    font-size: 0.95rem;
    font-weight: 300;
}
.badge {
    display: inline-block;
    background: rgba(99,102,241,0.15);
    color: #818cf8;
    border: 1px solid rgba(99,102,241,0.3);
    border-radius: 20px;
    padding: 3px 10px;
    font-size: 0.75rem;
    font-weight: 500;
    margin-right: 6px;
}
.badge.green { background: rgba(16,185,129,0.12); color: #34d399; border-color: rgba(16,185,129,0.3); }
.badge.amber { background: rgba(245,158,11,0.12); color: #fbbf24; border-color: rgba(245,158,11,0.3); }

/* Status cards */
.status-card {
    background: #1e2433;
    border: 1px solid #2a3050;
    border-radius: 12px;
    padding: 16px 20px;
    margin-bottom: 12px;
}
.status-dot {
    width: 8px; height: 8px;
    border-radius: 50%;
    display: inline-block;
    margin-right: 8px;
}
.dot-green { background: #10b981; box-shadow: 0 0 6px rgba(16,185,129,0.5); }
.dot-red   { background: #ef4444; }
.dot-amber { background: #f59e0b; box-shadow: 0 0 6px rgba(245,158,11,0.5); }

/* Live Meeting UI Enhancements */
.live-header {
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 20px;
    margin: 40px 0 10px;
}
.live-dot-3d {
    width: 60px;
    height: 60px;
    background: radial-gradient(circle at 30% 30%, #ff6b6b, #ee5253);
    border-radius: 50%;
    box-shadow: inset -5px -5px 15px rgba(0,0,0,0.4), 0 10px 20px rgba(238,82,83,0.3);
}
.live-title {
    font-family: 'DM Serif Display', serif;
    font-size: 3.5rem;
    color: #ffffff;
    margin: 0;
}
.live-subtitle {
    text-align: center;
    color: #94a3b8;
    font-size: 1.1rem;
    margin-bottom: 40px;
}
.action-card {
    background: #151921;
    border: 1px solid #2a3050;
    border-radius: 24px;
    padding: 40px;
    max-width: 800px;
    margin: 0 auto 30px;
    text-align: center;
}
.card-status-text {
    color: #e2e8f0;
    font-size: 1.2rem;
    margin-bottom: 30px;
}
.timer-row {
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 15px;
    margin-top: 30px;
    color: #94a3b8;
    font-size: 1.1rem;
    border-top: 1px solid #2a3050;
    padding-top: 20px;
}
.process-section {
    text-align: center;
    margin-top: 20px;
}
.process-help {
    color: #94a3b8;
    font-size: 0.9rem;
    margin-top: 10px;
}

/* Custom Button Styling */
div.stButton > button {
    border-radius: 12px;
    font-weight: 600;
}
/* Start button override */
div[data-testid="column"]:nth-child(1) div.stButton > button {
    background-color: #10b981 !important;
    color: white !important;
    border: none !important;
    height: 60px;
}
/* Stop button override */
div[data-testid="column"]:nth-child(2) div.stButton > button {
    background-color: #1a1f2e !important;
    color: #ef4444 !important;
    border: 1px solid #2a3050 !important;
    height: 60px;
}
/* Process button override */
.process-btn-container div.stButton > button {
    background-color: #1e2433 !important;
    color: #3b82f6 !important;
    border: 1px solid #2a3050 !important;
    padding: 15px 40px;
    font-size: 1.2rem;
}

/* Transcript box */
.transcript-box {
    background: #0d1117;
    border: 1px solid #2a3050;
    border-radius: 12px;
    padding: 20px;
    font-size: 0.92rem;
    line-height: 1.8;
    color: #cbd5e1;
    max-height: 400px;
    overflow-y: auto;
    white-space: pre-wrap;
    font-family: 'DM Sans', sans-serif;
}

/* Chat bubbles */
.chat-user {
    background: linear-gradient(135deg, #312e81, #1e1b4b);
    border: 1px solid rgba(99,102,241,0.3);
    border-radius: 18px 18px 4px 18px;
    padding: 12px 18px;
    margin: 8px 0;
    color: #e0e7ff;
    max-width: 80%;
    margin-left: auto;
    font-size: 0.93rem;
}
.chat-ai {
    background: #1e2433;
    border: 1px solid #2a3050;
    border-radius: 18px 18px 18px 4px;
    padding: 12px 18px;
    margin: 8px 0;
    color: #e2e8f0;
    max-width: 85%;
    font-size: 0.93rem;
    line-height: 1.7;
}
.chat-meta {
    font-size: 0.72rem;
    color: #64748b;
    margin-bottom: 4px;
}

/* Chunk progress */
.chunk-item {
    background: #161b27;
    border: 1px solid #2a3050;
    border-radius: 8px;
    padding: 8px 14px;
    margin: 4px 0;
    font-size: 0.82rem;
    color: #94a3b8;
}
.chunk-done { border-color: rgba(16,185,129,0.4); color: #86efac; }
.chunk-active { border-color: rgba(99,102,241,0.5); color: #c7d2fe; }

/* Provider pill */
.provider-pill {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    background: #1e2433;
    border: 1px solid #2a3050;
    border-radius: 20px;
    padding: 4px 12px;
    font-size: 0.8rem;
    color: #94a3b8;
}

.record-card {
    background: linear-gradient(180deg, rgba(15, 23, 42, 0.98), rgba(17, 24, 39, 0.98));
    border: 1px solid #334155;
    border-radius: 18px;
    padding: 18px;
    margin-bottom: 16px;
    box-shadow: 0 20px 50px rgba(15, 23, 42, 0.35);
}
.record-card-active {
    border-color: rgba(59, 130, 246, 0.85);
    box-shadow: 0 0 0 1px rgba(59, 130, 246, 0.25), 0 24px 60px rgba(59, 130, 246, 0.12);
}
.record-title {
    font-size: 1.1rem;
    font-weight: 800;
    color: #f8fafc;
}
.record-subtitle {
    color: #94a3b8;
    margin-top: 4px;
    font-size: 0.92rem;
}
.record-steps {
    margin-top: 12px;
    color: #cbd5e1;
    font-size: 0.92rem;
    line-height: 1.5;
}
.mic-spotlight {
    background: radial-gradient(circle at top, rgba(96,165,250,0.25), transparent 60%);
    border-radius: 16px;
    padding: 6px 0 0;
}
.mic-cta {
    background: linear-gradient(180deg, rgba(30,41,59,0.95), rgba(15,23,42,0.95));
    border: 1px solid #475569;
    border-radius: 18px;
    padding: 16px;
    margin-top: 10px;
    margin-bottom: 10px;
}
.mic-cta-title {
    font-size: 1rem;
    font-weight: 800;
    color: #f8fafc;
    margin-bottom: 6px;
}
.mic-cta-copy {
    color: #94a3b8;
    font-size: 0.92rem;
    margin-bottom: 12px;
}
</style>
""", unsafe_allow_html=True)

# ── Session state init ──────────────────────────────────────────────────────────
def init_state():
    defaults = {
        "transcript": "",
        "transcript_chunks": [],
        "chat_history": [],
        "audio_bytes": None,
        "processing": False,
        "meeting_title": "",
        "meeting_date": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "ai_provider": "auto",
        "selected_model": "",
        "transcript_segments": [],
        "diarization": None,
        "speakers": [],
        "speaker_map": {},
        "action_items": [],
        "mom": "",
        "pdf_bytes": None,
        "docx_bytes": None,
        "live_audio": None,
        "last_error": None,
        "last_traceback": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()

# ── Helper: detect available AI providers ──────────────────────────────────────
def detect_providers():
    providers = {}

    # Ollama (URL-based, user-configurable)
    ollama_url = st.secrets.get("OLLAMA_URL", "") or os.environ.get("OLLAMA_URL", "http://localhost:11434")
    try:
        import requests
        r = requests.get(f"{ollama_url}/api/tags", timeout=3)
        if r.status_code == 200:
            models = [m["name"] for m in r.json().get("models", [])]
            providers["ollama"] = {"url": ollama_url, "models": models or ["llama3", "mistral", "phi3"]}
    except Exception:
        pass

    # Groq
    groq_key = st.secrets.get("GROQ_API_KEY", "") or os.environ.get("GROQ_API_KEY", "")
    if groq_key:
        providers["groq"] = {
            "key": groq_key,
            "models": ["llama-3.1-8b-instant", "llama-3.3-70b-versatile", "openai/gpt-oss-20b", "openai/gpt-oss-120b"],
        }

    # Gemini
    gemini_key = st.secrets.get("GEMINI_API_KEY", "") or os.environ.get("GEMINI_API_KEY", "")
    if gemini_key:
        providers["gemini"] = {
            "key": gemini_key,
            "models": ["gemini-1.5-flash", "gemini-1.5-pro", "gemini-2.0-flash"],
        }

    return providers

# ── Helper: chunk audio file ────────────────────────────────────────────────────
def chunk_audio(audio_bytes: bytes, chunk_minutes: int = 5) -> list[bytes]:
    """Split audio into N-minute WAV chunks using ffmpeg when available."""
    ffmpeg = shutil.which("ffmpeg")
    if not ffmpeg:
        logger.warning("ffmpeg not found. Processing as single file.")
        st.info("ffmpeg is not available. Processing audio as a single file.")
        return [audio_bytes]

    input_path = None
    output_dir = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".bin", delete=False) as input_file:
            input_file.write(audio_bytes)
            input_path = input_file.name

        output_dir = tempfile.mkdtemp(prefix="meetingmind_chunks_")
        segment_pattern = os.path.join(output_dir, "chunk_%03d.wav")
        segment_seconds = max(30, int(chunk_minutes * 60))

        command = [
            ffmpeg, "-y", "-i", input_path, "-vn", "-acodec", "pcm_s16le",
            "-ar", "16000", "-ac", "1", "-f", "segment",
            "-segment_time", str(segment_seconds), "-reset_timestamps", "1",
            segment_pattern,
        ]

        logger.info(f"Chunking audio: {len(audio_bytes)} bytes into {chunk_minutes}m segments")
        result = subprocess.run(command, capture_output=True, text=True)
        if result.returncode != 0:
            logger.error(f"FFmpeg failed: {result.stderr}")
            raise RuntimeError(result.stderr.strip() or "ffmpeg chunking failed")

        chunks = []
        for name in sorted(os.listdir(output_dir)):
            if not name.lower().endswith(".wav"):
                continue
            with open(os.path.join(output_dir, name), "rb") as chunk_file:
                chunks.append(chunk_file.read())

        logger.info(f"Created {len(chunks)} audio chunks")
        if chunks:
            return chunks

        raise RuntimeError("ffmpeg produced no audio chunks")
    except Exception as e:
        logger.error(f"Chunking error: {str(e)}")
        st.info(f"Audio chunking unavailable ({e}). Processing as single file.")
        return [audio_bytes]
    finally:
        if input_path:
            try: os.unlink(input_path)
            except: pass
        if output_dir:
            try: shutil.rmtree(output_dir)
            except: pass

# ── Helper: transcribe with faster-whisper ─────────────────────────────────────
@st.cache_resource
def load_whisper(model_name: str = "base"):
    try:
        from faster_whisper import WhisperModel
        logger.info(f"Loading Whisper model: {model_name}")
        return WhisperModel(model_name, device="cpu", compute_type="int8")
    except Exception as e:
        logger.error(f"Failed to load Whisper: {str(e)}")
        return None


def transcribe_chunk(audio_bytes: bytes, language: str = "auto", model_name: str = "base") -> list:
    try:
        model = load_whisper(model_name)
        if model is None:
            return [{"text": "[Whisper failed to load]", "start": 0.0, "end": 0.0}]

        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
            f.write(audio_bytes)
            tmp_path = f.name

        kwargs = {}
        if language != "auto":
            kwargs["language"] = language

        segments, _ = model.transcribe(tmp_path, beam_size=5, **kwargs)
        out = []
        for seg in segments:
            out.append({"text": seg.text.strip(), "start": seg.start, "end": seg.end})

        try: os.unlink(tmp_path)
        except: pass
        return out
    except Exception as e:
        logger.error(f"Transcription error: {str(e)}")
        return [{"text": f"[Transcription error: {e}]", "start": 0.0, "end": 0.0}]


def map_segments_to_speakers(segments: list, diarization: list, unknown_label: str = "Unknown Speaker") -> list:
    if not diarization:
        for s in segments: s["speaker"] = unknown_label
        return segments

    for s in segments:
        s_start, s_end = s.get("start", 0.0), s.get("end", 0.0)
        best_sp, best_overlap = unknown_label, 0.0
        seg_dur = max(0.0, s_end - s_start)
        for d in diarization:
            overlap = max(0.0, min(s_end, d.get("end", 0.0)) - max(s_start, d.get("start", 0.0)))
            if overlap > best_overlap:
                best_overlap = overlap
                best_sp = d.get("speaker", unknown_label)

        if seg_dur > 0 and best_overlap / seg_dur < 0.2:
            s["speaker"] = unknown_label
        else:
            s["speaker"] = best_sp
    return segments


def extract_action_items(transcript: str) -> list:
    import re
    sents = re.split(r"(?<=[\.!?])\s+", transcript)
    action_keywords = ["action", "will", "assign", "todo", "due", "deadline", "deliver", "implement", "follow up", "owner"]
    actions = [s.strip() for s in sents if any(kw in s.lower() for kw in action_keywords) and len(s.strip()) > 10]
    return list(dict.fromkeys(actions))


def build_speaker_transcript(segments: list) -> str:
    lines = []
    for s in sorted(segments, key=lambda x: x.get("start", 0.0)):
        start = int(s.get("start", 0.0))
        ts = f"{start // 3600:02d}:{(start % 3600) // 60:02d}:{start % 60:02d}"
        sp = s.get("speaker") or "Unknown"
        name = st.session_state.speaker_map.get(sp, sp)
        lines.append(f"[{ts}] {name}: {s.get('text', '').replace('\n', ' ')}")
    return "\n".join(lines)


def run_diarization(audio_bytes: bytes):
    """Robust speaker diarization with strict dependency validation."""
    hf_token = st.secrets.get("HF_TOKEN", "") or os.environ.get("HF_TOKEN", "")
    if not hf_token:
        logger.warning("Diarization skipped: HF_TOKEN missing")
        return None, "HF_TOKEN not set. Add it to secrets to enable diarization."

    # 1. System Level Dependency Checks (Before Imports)
    if platform.system() == "Linux":
        if not shutil.which("ffmpeg"):
            logger.warning("Missing ffmpeg executable.")
            return None, "System Error: Missing ffmpeg. Diarization disabled on Streamlit Cloud without packages.txt."
        
        # Check libavutil explicitly to avoid torchcodec hard crash
        try:
            res = subprocess.run(["ldconfig", "-p"], capture_output=True, text=True)
            if "libavutil" not in res.stdout:
                logger.warning("Missing libavutil shared library.")
                return None, "System Error: Missing libavutil. Diarization disabled. Ensure packages.txt has ffmpeg."
        except Exception as e:
            logger.debug(f"ldconfig check failed: {e}")

    # 2. Python Level Dependency Checks
    try:
        import torch
        import torchaudio
        try:
            import torchcodec
        except ImportError as e:
            logger.warning(f"torchcodec missing or failed to load: {e}")
            return None, "Diarization disabled: torchcodec not installed or missing system libraries."
        except OSError as e:
            logger.warning(f"torchcodec OS Error (usually missing libavutil): {e}")
            return None, "Diarization disabled: missing system libraries required by torchcodec (libavutil)."

        import pyannote.audio
        from pyannote.audio import Pipeline
        from huggingface_hub import login
    except ImportError as e:
        logger.warning(f"Diarization dependencies missing: {e}")
        return None, "Diarization disabled: missing Python dependencies."
    except Exception as e:
        logger.warning(f"Unexpected dependency error: {e}")
        return None, "Diarization disabled due to environment issue."

    # 3. Execution
    try:
        logger.info(f"Initializing PyAnnote pipeline (pyannote.audio {pyannote.audio.__version__})")
        login(token=hf_token)

        pipeline = Pipeline.from_pretrained(
            "pyannote/speaker-diarization-3.1",
            use_auth_token=hf_token
        )
        
        if torch.cuda.is_available():
            pipeline.to(torch.device("cuda"))
            logger.info("Using CUDA for diarization")
        else:
            logger.info("Using CPU for diarization")

        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
            f.write(audio_bytes)
            wav_path = f.name

        logger.info(f"Running diarization on {len(audio_bytes)} bytes...")
        diar_result = pipeline(wav_path)
        
        diar_segments = []
        speakers = []
        for turn, track, speaker in diar_result.itertracks(yield_label=True):
            diar_segments.append({"start": turn.start, "end": turn.end, "speaker": speaker})
            speakers.append(speaker)
        
        try: os.unlink(wav_path)
        except: pass
        
        logger.info(f"Diarization complete: found {len(set(speakers))} speakers")
        return {"segments": diar_segments, "speakers": sorted(set(speakers))}, None

    except Exception as e:
        err_msg = str(e)
        logger.warning(f"Diarization execution failed: {err_msg}")
        logger.debug("Diarization traceback:\n" + traceback.format_exc())
        
        if "403" in err_msg or "gated" in err_msg.lower():
            return None, "Access Denied: Please accept terms for 'pyannote/speaker-diarization-3.1' on Hugging Face."
        elif "token" in err_msg.lower():
            return None, "Invalid HF_TOKEN. Please check your credentials."
        elif "libavutil" in err_msg or "ffmpeg" in err_msg.lower() or "torchcodec" in err_msg.lower():
            return None, "System Error: Missing audio libraries (FFmpeg/libavutil). Diarization requires system-level dependencies."
        return None, "Diarization failed during processing (likely Out of Memory). Check logs."



def process_audio_bytes(audio_bytes: bytes, language: str, chunk_minutes: int, enable_diarization: bool = True) -> None:
    st.session_state.processing = True
    logger.info("Starting audio processing pipeline...")
    
    with st.spinner("Chunking audio..."):
        chunks = chunk_audio(audio_bytes, chunk_minutes=chunk_minutes)

    total_segments = []
    progress = st.progress(0, text="Transcribing...")
    for i, chunk in enumerate(chunks):
        segs = transcribe_chunk(chunk, language=language)
        total_segments.extend(segs)
        progress.progress((i + 1) / len(chunks), text=f"Transcribing chunk {i+1}/{len(chunks)}...")

    st.session_state.transcript_segments = total_segments
    st.session_state.transcript = "\n\n".join(s.get("text", "") for s in total_segments)
    st.session_state.action_items = extract_action_items(st.session_state.transcript)

    if enable_diarization:
        with st.spinner("Running speaker diarization..."):
            diar_data, error = run_diarization(audio_bytes)
            if diar_data:
                st.session_state.diarization = diar_data["segments"]
                st.session_state.speakers = diar_data["speakers"]
                st.session_state.speaker_map = {s: s for s in diar_data["speakers"]}
                st.session_state.transcript_segments = map_segments_to_speakers(
                    st.session_state.transcript_segments,
                    st.session_state.diarization
                )
            else:
                st.warning(f"Speaker Diarization Fallback: {error}")
                logger.warning(f"Diarization fallback triggered: {error}")

    st.session_state.processing = False
    logger.info("Audio processing complete.")


# ── AI Integration ─────────────────────────────────────────────────────────────
def ask_ai(question: str, transcript: str, history: list, provider: str, model: str, providers: dict) -> str:
    system_prompt = f"You are MeetingMind, an expert meeting analyst AI. Answer based ONLY on the transcript:\n\n{transcript[:12000]}"
    messages = [{"role": "system", "content": system_prompt}]
    for turn in history[-10:]: messages.append({"role": turn["role"], "content": turn["content"]})
    messages.append({"role": "user", "content": question})

    try:
        if provider == "ollama":
            import requests
            r = requests.post(f"{providers['ollama']['url']}/api/chat", json={"model": model, "messages": messages, "stream": False}, timeout=120)
            r.raise_for_status()
            return r.json()["message"]["content"]
        elif provider == "groq":
            from groq import Groq
            client = Groq(api_key=providers["groq"]["key"])
            resp = client.chat.completions.create(model=model, messages=messages, temperature=0.3)
            return resp.choices[0].message.content
        elif provider == "gemini":
            import google.generativeai as genai
            genai.configure(api_key=providers["gemini"]["key"])
            gmodel = genai.GenerativeModel(model_name=model, system_instruction=system_prompt)
            chat = gmodel.start_chat(history=[{"role": "user" if t["role"] == "user" else "model", "parts": [t["content"]]} for t in history[-10:]])
            return chat.send_message(question).text
    except Exception as e:
        logger.error(f"AI Provider error ({provider}): {str(e)}")
        return f"AI Error: {e}"
    return "No AI provider available."


@st.cache_data(ttl=3600)
def generate_mom_cached(title: str, date: str, duration: str, speakers: str, transcript: str, actions: list, provider: str, model: str, providers: dict) -> str:
    prompt = MOM_PROMPT.format(title=title, date=date, duration=duration, speakers=speakers, transcript=transcript)
    if actions: prompt += "\n\nEXTRACTED_ACTIONS:\n" + "\n".join(f"- {a}" for a in actions)
    return ask_ai(prompt, transcript, [], provider, model, providers)

MOM_PROMPT = """# Meeting Intelligence Report
**Meeting Title:** {title} | **Date:** {date} | **Duration:** {duration}
**Attendees:** {speakers}

## 1. Executive Summary
[Provide a concise 2-3 sentence overview of the meeting's primary purpose and outcome]

## 2. Key Discussion Points
[Bullet points of the main topics discussed, grouped logically]

## 3. Decisions Made
[Numbered list of all final decisions agreed upon during the meeting]

## 4. Risks & Issues
[Bullet points of any blockers, risks, or concerns raised by participants]

## 5. Action Items
| # | Action | Owner | Deadline |
|---|--------|-------|----------|
[Extract every task or action item mentioned. Assign to the speaker who agreed to it, or "Unassigned". Estimate deadlines if mentioned.]

## 6. Next Steps
[What happens immediately after this meeting]

TRANSCRIPT REFERENCE:
{transcript}"""

def export_pdf(mom_text: str, logo_bytes: bytes | None = None) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        c.drawString(20, 800, "Minutes of Meeting")
        t = c.beginText(20, 780)
        for line in mom_text.splitlines(): t.textLine(line)
        c.drawText(t)
        c.showPage()
        c.save()
        return buf.getvalue()
    except Exception as e:
        st.error(f"PDF error: {e}")
        return b""

def export_docx(mom_text: str) -> bytes:
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Minutes of Meeting", 0)
        for line in mom_text.splitlines(): doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        st.error(f"DOCX error: {e}")
        return b""

# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("### ⚙️ Configuration")
    providers = detect_providers()
    
    # Provider status
    st.markdown("**AI Providers**")
    for p, info in providers.items():
        st.markdown(f'<div class="status-card" style="padding:10px 14px;"><span class="status-dot dot-green"></span> {p.capitalize()}</div>', unsafe_allow_html=True)

    if providers:
        provider_choice = st.selectbox("Provider", list(providers.keys()))
        model_choice = st.selectbox("Model", providers[provider_choice].get("models", []))
        st.session_state.ai_provider, st.session_state.selected_model = provider_choice, model_choice

    st.divider()
    whisper_lang = st.selectbox("Language", ["auto", "en", "ta", "hi", "fr", "de", "es", "zh"])
    chunk_size = st.slider("Chunk size (minutes)", 2, 10, 5)
    
    st.divider()
    st.session_state.meeting_title = st.text_input("Title", st.session_state.meeting_title)
    st.session_state.meeting_date = st.text_input("Date", st.session_state.meeting_date)

    if st.button("🗑 Clear Session", use_container_width=True):
        st.session_state.clear()
        st.rerun()

    # Debug Panel
    with st.expander("🛠 System Diagnostics"):
        if st.button("Run Diagnostics"):
            with st.spinner("Analyzing environment..."):
                diag = get_system_diagnostics()
                st.json(diag)
        
        if st.session_state.last_error:
            st.error(f"Last Error: {st.session_state.last_error}")
            with st.expander("View Traceback"):
                st.code(st.session_state.last_traceback)
        
        if LOG_FILE.exists():
            st.download_button("Download Logs", data=LOG_FILE.read_text(), file_name="app.log")

# ══════════════════════════════════════════════════════════════════════════════
# MAIN UI
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="main-header"><h1>🎙 MeetingMind</h1><p>AI-powered meeting transcription & analysis</p></div>', unsafe_allow_html=True)

tab_live, tab_upload, tab_transcript, tab_mom = st.tabs([
    "🔴 Live",
    "📁 Upload",
    "📋 Transcript",
    "📄 Minutes",
])

# -----------------------------
# TAB: Live Meeting
# -----------------------------
with tab_live:
    st.markdown("""
        <div class="live-header" style="margin-top: 20px;">
            <div class="live-dot-3d"></div>
            <h1 class="live-title">Live Meeting Mode</h1>
        </div>
        <p class="live-subtitle">Capture browser audio from Google Meet, Teams, or Zoom.</p>
    """, unsafe_allow_html=True)

    recorder_data = render_tab_recorder()
    
    if recorder_data and recorder_data.get("action") == "process":
        if "audio_base64" in recorder_data:
            audio_b64 = recorder_data["audio_base64"]
            duration = recorder_data["duration"]
            audio_bytes = base64.b64decode(audio_b64)
            
            # Prevent double-processing if component re-renders
            if st.session_state.get("last_processed_duration") != duration:
                st.session_state.live_audio_data = audio_bytes
                st.session_state.last_processed_duration = duration
                
                st.success(f"✅ Audio captured successfully. Duration: {duration}")
                
                # Immediately trigger processing
                process_audio_bytes(audio_bytes, whisper_lang, chunk_size)
                
                # Post-processing Meeting Intelligence Extraction
                st.success("Audio processed! Extracting Meeting Intelligence...")
                provider = st.session_state.ai_provider if st.session_state.ai_provider != "auto" else provider_choice or next(iter(providers), None)
                model = st.session_state.selected_model or model_choice
                
                if provider and st.session_state.transcript:
                    with st.spinner("Generating AI Summary, Actions, and Decisions..."):
                        speakers_list = ", ".join(st.session_state.speakers) if st.session_state.get("speakers") else "TBD"
                        tagged_transcript = build_speaker_transcript(st.session_state.transcript_segments) if st.session_state.transcript_segments else st.session_state.transcript
                        
                        mom_text = generate_mom_cached(
                            title=st.session_state.get("meeting_title", "Live Meeting"),
                            date=st.session_state.get("meeting_date", ""),
                            duration=duration,
                            speakers=speakers_list,
                            transcript=tagged_transcript,
                            actions=st.session_state.get("action_items", []),
                            provider=provider,
                            model=model,
                            providers=providers,
                        )
                        st.session_state.mom = mom_text
                    st.success("Meeting Intelligence generated! Check the Minutes tab.")
                else:
                    st.info("Check the Transcript tab for results. Configure an AI Provider in the sidebar to generate Minutes.")
            else:
                st.info(f"Processing complete for session: {duration}")

with tab_upload:
    uploaded = st.file_uploader("Upload audio", type=["mp3", "wav", "m4a", "ogg"])
    if uploaded:
        st.audio(uploaded)
        if st.button("Process Uploaded Audio"):
            process_audio_bytes(uploaded.read(), whisper_lang, chunk_size)

with tab_transcript:
    if not st.session_state.transcript:
        st.info("No transcript yet.")
    else:
        st.download_button("Download .txt", st.session_state.transcript, "transcript.txt")
        if st.session_state.transcript_segments:
            for s in st.session_state.transcript_segments:
                name = st.session_state.speaker_map.get(s.get("speaker"), s.get("speaker", "Unknown"))
                st.markdown(f"**{name}**: {s.get('text')}")
        else:
            st.text_area("Transcript", st.session_state.transcript, height=400)

with tab_mom:
    if not st.session_state.transcript:
        st.info("No transcript yet.")
    else:
        if st.button("Generate MoM"):
            st.session_state.mom = generate_mom_cached(st.session_state.meeting_title, st.session_state.meeting_date, "TBD", ", ".join(st.session_state.speakers), st.session_state.transcript, st.session_state.action_items, st.session_state.ai_provider, st.session_state.selected_model, providers)
        if st.session_state.mom:
            st.text_area("MoM", st.session_state.mom, height=400)
            st.download_button("Download TXT", st.session_state.mom, "mom.txt")
