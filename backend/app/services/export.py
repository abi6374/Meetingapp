import io
import logging
import json
from datetime import datetime
from typing import Any, Dict

logger = logging.getLogger(__name__)

def _clean_json_string(text: str) -> str:
    """Removes markdown fences and whitespace from AI-generated JSON."""
    text = text.strip()
    if text.startswith("```"):
        text = text.replace("```json", "", 1).replace("```", "", 1)
        if text.endswith("```"):
            text = text[:-3]
    return text.strip()

def _parse_mom_json(text: str) -> Dict[str, Any] | None:
    """Attempts to parse MOM text as structured JSON."""
    try:
        clean_text = _clean_json_string(text)
        return json.loads(clean_text)
    except Exception:
        return None

def export_pdf(mom_text: str) -> bytes:
    """Generate a formal Minutes of Meeting (MOM) PDF using ReportLab Platypus."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
        
        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, 
            pagesize=A4,
            rightMargin=40, leftMargin=40,
            topMargin=40, bottomMargin=40
        )
        
        styles = getSampleStyleSheet()
        
        # ── CUSTOM FORMAL STYLES ──────────────────────────────────────────────
        mom_title_style = ParagraphStyle(
            'MOMTitle',
            parent=styles['Heading1'],
            fontSize=26,
            textColor=colors.HexColor("#0f172a"),
            alignment=TA_CENTER,
            spaceAfter=10,
            fontName='Helvetica-Bold'
        )
        
        doc_info_style = ParagraphStyle(
            'DocInfo',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor("#64748b"),
            alignment=TA_CENTER,
            spaceAfter=20
        )
        
        section_header_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor("#1e293b"),
            spaceBefore=15,
            spaceAfter=8,
            fontName='Helvetica-Bold',
            borderPadding=2,
            underlineWidth=1
        )

        agenda_topic_style = ParagraphStyle(
            'AgendaTopic',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor("#334155"),
            fontName='Helvetica-Bold',
            spaceBefore=10,
            spaceAfter=4
        )
        
        body_text_style = ParagraphStyle(
            'BodyText',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor("#334155"),
            spaceAfter=6,
            leading=14,
            alignment=TA_LEFT
        )

        decision_style = ParagraphStyle(
            'DecisionText',
            parent=body_text_style,
            leftIndent=20,
            textColor=colors.HexColor("#059669"),
            fontName='Helvetica-BoldOblique'
        )

        label_style = ParagraphStyle(
            'LabelStyle',
            parent=body_text_style,
            fontName='Helvetica-Bold',
            textColor=colors.HexColor("#475569")
        )

        elements = []
        
        # Parse data
        mom_data = _parse_mom_json(mom_text)
        
        if mom_data:
            header = mom_data.get('header', {})
            
            # 1. FORMAL HEADER
            elements.append(Paragraph("MINUTES OF MEETING", mom_title_style))
            elements.append(Paragraph(header.get('title', 'Untitled Meeting').upper(), ParagraphStyle('SubTitle', parent=mom_title_style, fontSize=16)))
            elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceBefore=10, spaceAfter=10))
            
            # 2. MEETING INFO GRID
            info_table_data = [
                [Paragraph("<b>Date:</b>", label_style), Paragraph(header.get('date') or "N/A", body_text_style), 
                 Paragraph("<b>Location:</b>", label_style), Paragraph(header.get('location') or "Virtual", body_text_style)],
                [Paragraph("<b>Time:</b>", label_style), Paragraph(header.get('time') or "N/A", body_text_style), 
                 Paragraph("<b>Facilitator:</b>", label_style), Paragraph(header.get('facilitator') or "TBD", body_text_style)]
            ]
            info_table = Table(info_table_data, colWidths=[80, 180, 80, 180])
            info_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
            ]))
            elements.append(info_table)
            elements.append(Spacer(1, 10))

            # 3. ATTENDANCE
            elements.append(Paragraph("1. ATTENDANCE", section_header_style))
            attendance = mom_data.get('attendance', {})
            present = attendance.get('present', [])
            if present:
                elements.append(Paragraph(f"<b>Present:</b> {', '.join(present)}", body_text_style))
            else:
                elements.append(Paragraph("<b>Present:</b> [No participants identified]", body_text_style))
            
            absent = attendance.get('absent', [])
            if absent:
                elements.append(Paragraph(f"<b>Absent:</b> {', '.join(absent)}", body_text_style))

            # 4. AGENDA & DISCUSSION
            elements.append(Paragraph("2. AGENDA ITEMS & DISCUSSION", section_header_style))
            agenda_items = mom_data.get('agenda_items', [])
            if agenda_items:
                for idx, item in enumerate(agenda_items, 1):
                    elements.append(Paragraph(f"{idx}. {item.get('topic', 'Topic')}", agenda_topic_style))
                    elements.append(Paragraph(item.get('discussion', ''), body_text_style))
                    
                    decisions = item.get('decisions', [])
                    if decisions:
                        for dec in decisions:
                            elements.append(Paragraph(f"<b>DECISION:</b> {dec}", decision_style))
                    elements.append(Spacer(1, 5))
            else:
                elements.append(Paragraph("No specific agenda items recorded.", body_text_style))

            # 5. CONSOLIDATED ACTION ITEMS
            elements.append(Paragraph("3. ACTION ITEMS", section_header_style))
            
            # Flatten all action items from agenda
            all_actions = []
            for item in agenda_items:
                for action in item.get('action_items', []):
                    all_actions.append(action)
            
            if all_actions:
                table_data = [['#', 'Action/Task', 'Owner', 'Deadline']]
                for i, action in enumerate(all_actions, 1):
                    table_data.append([
                        str(i),
                        Paragraph(action.get('task', 'N/A'), body_text_style),
                        Paragraph(action.get('owner', 'TBD'), body_text_style),
                        Paragraph(action.get('deadline') or "N/A", body_text_style)
                    ])
                
                t = Table(table_data, colWidths=[25, 265, 110, 110])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#f8fafc")),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor("#1e293b")),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0,0), (-1,0), 10),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('TOPPADDING', (0,1), (-1,-1), 8),
                    ('BOTTOMPADDING', (0,1), (-1,-1), 8),
                ]))
                elements.append(t)
            else:
                elements.append(Paragraph("No immediate action items identified.", body_text_style))

            # 6. OTHER BUSINESS & NEXT STEPS
            elements.append(Paragraph("4. OTHER BUSINESS", section_header_style))
            gen_notes = mom_data.get('general_notes')
            if gen_notes:
                elements.append(Paragraph(gen_notes, body_text_style))
            
            next_sync = mom_data.get('next_sync')
            if next_sync:
                elements.append(Paragraph(f"<b>NEXT MEETING:</b> {next_sync}", agenda_topic_style))
            
            elements.append(Spacer(1, 40))
            elements.append(Paragraph(f"Report generated by MeetingMind AI on {datetime.now().strftime('%Y-%m-%d %H:%M')}", doc_info_style))

        else:
            # FALLBACK (PLAIN TEXT)
            elements.append(Paragraph("Meeting Intelligence Report", mom_title_style))
            for line in mom_text.splitlines():
                elements.append(Paragraph(line, body_text_style))

        def add_footer(canvas, doc):
            canvas.saveState()
            canvas.setFont('Helvetica-Oblique', 8)
            canvas.setStrokeColor(colors.HexColor("#e2e8f0"))
            canvas.line(40, 30, A4[0]-40, 30)
            canvas.drawString(40, 20, f"Confidential - MeetingMind Intelligence")
            canvas.drawRightString(A4[0]-40, 20, f"Page {doc.page}")
            canvas.restoreState()

        doc.build(elements, onFirstPage=add_footer, onLaterPages=add_footer)
        return buf.getvalue()
    except Exception as e:
        logger.error(f"Formal PDF export failed: {e}")
        return _fallback_export_pdf(mom_text)

def _fallback_export_pdf(mom_text: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        c.setFont("Helvetica-Bold", 16)
        c.drawString(40, 800, "Meeting Intelligence Report")
        c.setFont("Helvetica", 10)
        t = c.beginText(40, 780)
        for line in mom_text.splitlines():
            t.textLine(line)
        c.drawText(t)
        c.showPage()
        c.save()
        return buf.getvalue()
    except Exception as e:
        logger.error(f"Fallback PDF export failed: {e}")
        return b""

def export_docx(mom_text: str) -> bytes:
    """Generate a formal DOCX Minutes of Meeting."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        mom_data = _parse_mom_json(mom_text)
        
        if mom_data:
            header = mom_data.get('header', {})
            
            # Title
            title = doc.add_heading("MINUTES OF MEETING", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            meeting_title = doc.add_heading(header.get('title', 'Untitled Meeting').upper(), level=1)
            meeting_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Header Grid
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = f"Date: {header.get('date') or 'N/A'}"
            table.cell(0, 1).text = f"Location: {header.get('location') or 'Virtual'}"
            table.cell(1, 0).text = f"Time: {header.get('time') or 'N/A'}"
            table.cell(1, 1).text = f"Facilitator: {header.get('facilitator') or 'TBD'}"
            
            # 1. Attendance
            doc.add_heading("1. ATTENDANCE", level=2)
            attendance = mom_data.get('attendance', {})
            doc.add_paragraph(f"Present: {', '.join(attendance.get('present', []))}")
            if attendance.get('absent'):
                doc.add_paragraph(f"Absent: {', '.join(attendance.get('absent', []))}")
            
            # 2. Agenda
            doc.add_heading("2. AGENDA ITEMS & DISCUSSION", level=2)
            for item in mom_data.get('agenda_items', []):
                doc.add_heading(item.get('topic', 'Topic'), level=3)
                doc.add_paragraph(item.get('discussion', ''))
                for dec in item.get('decisions', []):
                    p = doc.add_paragraph(f"DECISION: {dec}")
                    p.style.font.bold = True
            
            # 3. Action Items
            doc.add_heading("3. ACTION ITEMS", level=2)
            actions = []
            for item in mom_data.get('agenda_items', []):
                actions.extend(item.get('action_items', []))
            
            if actions:
                a_table = doc.add_table(rows=1, cols=4)
                a_table.style = 'Table Grid'
                hdr = a_table.rows[0].cells
                hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '#', 'Task', 'Owner', 'Deadline'
                for i, a in enumerate(actions, 1):
                    row = a_table.add_row().cells
                    row[0].text, row[1].text, row[2].text, row[3].text = str(i), a.get('task', 'N/A'), a.get('owner', 'TBD'), a.get('deadline', 'N/A')
            
            # 4. Other
            doc.add_heading("4. OTHER BUSINESS", level=2)
            doc.add_paragraph(mom_data.get('general_notes', 'N/A'))
            if mom_data.get('next_sync'):
                doc.add_paragraph(f"Next Meeting: {mom_data.get('next_sync')}")

        else:
            doc.add_heading("Meeting Intelligence Report", 0)
            for line in mom_text.splitlines():
                doc.add_paragraph(line)
        
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        logger.error(f"Formal DOCX export failed: {e}")
        return b""
